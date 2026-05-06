"""
生成导师汇报 Word：CcVault 汽车法规知识库 (2026-04-22)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------- #
#  样式工具
# ---------------------------------------------------------------- #
def set_cn_font(run, font_name: str = "微软雅黑", size: float = 10.5):
    run.font.name = font_name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_cn_font(run, "微软雅黑", 18 - level * 2)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text: str, bold: bool = False, size: float = 10.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, "微软雅黑", size)
    if bold:
        run.font.bold = True
    return p


def add_bullet(doc, text: str, size: float = 10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_cn_font(run, "微软雅黑", size)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]]):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = True

    # 表头
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_cn_font(run, "微软雅黑", 10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # 底色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3A5F")
        tcPr.append(shd)

    # 数据行
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_cn_font(run, "微软雅黑", 10)
    return tbl


# ---------------------------------------------------------------- #
#  内容生成
# ---------------------------------------------------------------- #
def build_doc() -> Document:
    doc = Document()

    # 全局页边距
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CcVault 全球汽车法规知识库")
    set_cn_font(run, "微软雅黑", 22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("研究进展汇报 · 2026 年 4 月 22 日")
    set_cn_font(run, "微软雅黑", 13)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # 空行

    # ================================================================
    # 一、项目定位
    # ================================================================
    add_heading(doc, "一、项目定位：做了什么", level=1)

    add_para(
        doc,
        "CcVault 是一套面向毕业设计的全球汽车法规结构化知识库。"
        "项目把原本散落在 1537 份 PDF / Word / Excel 中的法规文本（合计 1.81 GB，"
        "涵盖中国 GB/GB/T、联合国 ECE、欧盟指令、美国 FMVSS、日韩 JIS/KMVSS、"
        "ISO/SAE 等 17 个区域），通过 OCR、结构化抽取、跨模型交叉审校等自动化流水线，"
        "转化为 1 415 条标准化的 Obsidian Markdown 条目，并在其上搭建语义检索、"
        "GraphRAG 社区综述、ReAct Agent 对话问答三层应用能力。",
        size=11,
    )

    add_para(doc, "项目的四个定位：", bold=True)
    add_bullet(doc, "个人工具箱：工作查询（例：GB 11551-2014 HIC 限值）秒级响应。")
    add_bullet(doc, "论文脚手架：系统对比多区法规版本演进、跨区差异、监管理念。")
    add_bullet(doc, "职业竞争力：等价于跨国车企工程师的法规结构化能力。")
    add_bullet(doc, "持续学习：关系图随新文档入库不断生长。")

    # ================================================================
    # 二、核心功能
    # ================================================================
    add_heading(doc, "二、知识库的核心功能", level=1)

    add_table(
        doc,
        headers=["功能模块", "能力说明", "当前规模"],
        rows=[
            ["结构化 Note", "每条法规抽成 YAML FrontMatter + 正文，字段统一", "1 415 条"],
            ["跨区域等价映射", "GB ↔ ECE / EU / ISO 对标关系自动维护", "169 条"],
            ["主题聚类", "按 brakes / lighting / EMC 等 37 个技术主题索引", "37 主题"],
            ["BM25 关键词检索", "jieba 中文分词 + BM25，毫秒级返回", "全库 1.6 MB 索引"],
            ["GraphRAG 层级检索", "Louvain 聚类生成 33 个社区综述，支持跨法规推理", "33 社区"],
            ["ReAct Agent 对话", "自然语言问答，支持 12 个工具（统计 / 检索 / 沿革）", "CLI+HTTP+MCP 三接入"],
            ["Dataview 活面板", "Obsidian 端查询体检、低置信度审核队列等", "10 面板"],
            ["版本沿革链", "自动维护 supersedes / superseded_by 双向链", "89 条 superseded"],
        ],
    )

    # ================================================================
    # 三、流水线六阶段
    # ================================================================
    add_heading(doc, "三、流水线：六阶段实现", level=1)

    add_para(
        doc,
        "知识库的数据生产由一条可重跑的 Python 流水线驱动，分六个阶段，"
        "每个阶段都有独立的 CLI 入口和中间产物，支持断点续跑与局部重建。",
    )

    add_table(
        doc,
        headers=["阶段", "任务", "关键技术 / 模型", "产出"],
        rows=[
            [
                "Stage 1\nOCR 分层",
                "把 PDF 还原为 Markdown，按 PDF 类型自动分流",
                "pdfplumber（电子版）\n百度云 OCR（扫描件）\nMinerU 云 API（复杂表格 / 公式 / 图像）",
                "_staging / *.md",
            ],
            [
                "Stage 2\n结构化抽取",
                "从 Markdown 抽 FrontMatter（reg_id、publication_date、scope 等）+ 章节结构",
                "DeepSeek V3（¥300-350 批量）",
                "01_Wiki/regulations/*.md",
            ],
            [
                "Stage 3\nCross-check 审校",
                "对每条 note 做字段一致性、引用完整性、摘要准确性三项交叉审核",
                "Claude Sonnet 4.6 Batch API",
                "cross_check_* 字段 +\nneeds-review 标签",
            ],
            [
                "Stage 4\n跨区等价判定",
                "识别 GB 4785 ↔ ECE R48 这类对标关系，5 种关系类型",
                "Claude Opus 4.7 Batch API",
                "03_Equivalence / *.md",
            ],
            [
                "Stage 5\n主题聚类 + GraphRAG",
                "规则分类到 37 主题 + 基于引用图的 Louvain 社区检测 + LLM 社区综述",
                "jieba + networkx + python-louvain\nClaude Opus 4.7（社区综述）",
                "04_Topics / communities/\n33 社区摘要",
            ],
            [
                "Stage 6\n检索与 Agent",
                "构建 BM25 索引 + ReAct Agent 12 工具 + HTTP/MCP 接入",
                "rank_bm25 + httpx + MCP 协议\nClaude Opus 4.7（问答）",
                "bm25_index.pkl\n_agent_{chat,server,mcp}.py",
            ],
        ],
    )

    add_para(doc, "运行入口：", bold=True)
    add_bullet(doc, "全流程：python 99_SystemScripts/auto_reg_index/ingest.py --phase 2 --batch-all")
    add_bullet(doc, "日维护：python _daily_maintenance.py")
    add_bullet(doc, "新 PDF 入库：/ingest workflow（.windsurf/workflows/ingest.md）")

    # ================================================================
    # 四、当前成果量化
    # ================================================================
    add_heading(doc, "四、当前成果（截至 2026-04-22）", level=1)

    add_table(
        doc,
        headers=["指标", "数值"],
        rows=[
            ["结构化 Note 数", "1 415 条（17 个区域，35+ 主题）"],
            ["数据质量（cross_check）", "high 1 016 · medium 355 · low 35（87.4% verified）"],
            ["跨区域等价映射", "169 条"],
            ["GraphRAG 社区", "33 个（Louvain 聚类 + Sonnet 4.6 综述）"],
            ["原文 OCR 升级（MinerU 云）", "572 / 1 444 PDF 已升级（39%），含表格 1 442 / 公式 806 / 图像 1 561"],
            ["索引体积", "BM25 1.6 MB · GraphRAG 图谱 .stage5/graph.json"],
            ["总 API 成本", "¥102 / ≈ $14.33（低于预算 ¥1 000）"],
            ["代码量", "Python 流水线 ≈ 15 000 行，30+ 独立脚本"],
            ["Agent 工具", "12 个（统计 / BM25 / GraphRAG / 等价 / 沿革 / 审核等）"],
        ],
    )

    # ================================================================
    # 五、今日工作摘要
    # ================================================================
    add_heading(doc, "五、本日（2026-04-22）工作重点", level=1)

    add_para(doc, "重点推进了「原文 OCR 二次升级」与「数据质量修复」两块。", bold=True)

    add_bullet(doc, "接入 MinerU 云 API，对 572 份重要法规 PDF 执行二次高精度 OCR，将表格 / 公式 / 图像三类元素回填到 note 正文。")
    add_bullet(doc, "修复合并脚本 5 处 bug：FM 空行、空图像占位、表格计数虚高、SSL polling 崩溃、合并计数对齐。")
    add_bullet(doc, "新增 3 个自愈机制：上传失败重试（3 次指数退避）、SSL EOF 恢复、小文件优先调度。")
    add_bullet(doc, "MinerU 单日上传成功率从 10-40% 提升到 100%。")
    add_bullet(doc, "完成 dedupe（15 条重复 note 合并）、manifest 同步、BM25 索引重建、合并 QC 回填。")

    # ================================================================
    # 六、未完成事项
    # ================================================================
    add_heading(doc, "六、尚未完成 / 待推进", level=1)

    add_para(doc, "按优先级分为 4 档：", bold=True)

    add_table(
        doc,
        headers=["优先级", "任务", "预估工时", "说明"],
        rows=[
            [
                "P0\n本月",
                "MinerU 全量跑完",
                "3-5 天",
                "剩余 872 份 PDF（约 60%）待升级，受 MinerU 服务器拥塞影响，需分日跑黄金窗口",
            ],
            [
                "P0\n本月",
                "Dense Vector 检索接入",
                "2-3 周",
                "当前只有 BM25，计划接入 bge-m3 做 Hybrid（BM25 + 向量 + rerank）",
            ],
            [
                "P1\n本季度",
                "GraphRAG 升级",
                "2-4 周",
                "Louvain → Leiden、一轮综述 → 迭代精炼、加上 entity extraction",
            ],
            [
                "P1\n本季度",
                "段落级检索",
                "1-2 周",
                "长法规（如 ECE R13 200 KB）按章节 / Annex 切块，提升命中精度",
            ],
            [
                "P1\n本季度",
                "Agent Write-tools",
                "3-5 周",
                "当前 Agent 只读；计划加自动 /add_note /fix_classification 等写操作",
            ],
            [
                "P2\n半年内",
                "Web / Mobile 前端",
                "4-8 周",
                "脱离 Obsidian，面向多用户的法规检索 Web 应用",
            ],
            [
                "P2\n半年内",
                "条款级差异分析",
                "4-8 周",
                "GB 4785-2019 vs ECE R48 Rev6 的条款级语义对比，辅助合规性检查",
            ],
            [
                "P3\n长期",
                "多人协作 / 权限管理",
                "8-12 周",
                "从单人知识库升级为团队平台，涉及账号、权限、审计",
            ],
        ],
    )

    add_para(doc, "数据质量层面待修的小问题：", bold=True)
    add_bullet(doc, "35 条低置信度（cross_check=low）note 待人工复核。")
    add_bullet(doc, "9 条 note 的 cross_check 字段为空，需补跑。")
    add_bullet(doc, "2 条 LLM 生成的图像占位符（如「图 4 描述」）需替换为真实 MinerU 抽取的图像。")

    # ================================================================
    # 七、技术栈总览
    # ================================================================
    add_heading(doc, "七、技术栈总览", level=1)

    add_table(
        doc,
        headers=["层次", "技术选型"],
        rows=[
            ["数据层", "Obsidian Vault + Markdown + YAML FrontMatter"],
            ["OCR", "pdfplumber（电子 PDF） / 百度云 OCR（扫描件） / MinerU 云 API（复杂文档）"],
            ["LLM 分工", "DeepSeek V3（批量抽取） / Claude Sonnet 4.6（审校、主题、社区综述） / Claude Opus 4.7（跨区判定、问答）"],
            ["调用层", "Anthropic Batch API（省 50% 成本） + httpx + 自研 SSL 自愈客户端"],
            ["检索层", "rank_bm25 + jieba 分词（1.6 MB pickle 索引）"],
            ["图谱层", "networkx + python-louvain（Louvain 聚类）"],
            ["Agent 层", "自研 ReAct 框架 + 12 tools · CLI + HTTP（OpenAI 兼容）+ MCP（Claude Desktop / Cursor / Cascade）"],
            ["前端", "Obsidian + Dataview + Templater + Smart Composer + Copilot（本地 Agent 接入）"],
            ["CI / 运维", "Python 流水线 + .windsurf/workflows（5 个命令式 workflow）+ 每日维护脚本"],
        ],
    )

    # ================================================================
    # 结束
    # ================================================================
    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = end.add_run("—— 汇报结束 ——")
    set_cn_font(run, "微软雅黑", 10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


if __name__ == "__main__":
    out = Path(r"D:\4月22日汇报.docx")
    doc = build_doc()
    doc.save(out)
    print(f"✓ 已生成: {out}")
    print(f"  文件大小: {out.stat().st_size / 1024:.1f} KB")
